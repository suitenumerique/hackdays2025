from bs4 import BeautifulSoup
from docx import Document

def parse_html_to_docx(content_html: str) -> Document:
    soup = BeautifulSoup(content_html, 'html.parser')
    doc = Document()
    doc.add_heading('Résumé généré', 0)

    for element in soup.contents:
        if element.name == 'h3':
            doc.add_paragraph(element.get_text(strip=True), style='Heading 2')
        elif element.name == 'b':
            p = doc.add_paragraph()
            p.add_run(element.get_text(strip=True)).bold = True
        elif element.name == 'br':
            doc.add_paragraph()  # ligne vide
        elif element.name is None:
            if element.strip():
                doc.add_paragraph(element.strip())
        else:
            doc.add_paragraph(element.get_text(strip=True))

    return doc
